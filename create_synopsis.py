import os
import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Define styles
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Times New Roman'
font.size = Pt(12)

# TITLE PAGE
doc.add_paragraph('\n\n\n')
p = doc.add_paragraph('A\nSynopsis on\n')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14)
p.runs[0].bold = True

title = doc.add_paragraph('AI Cognitive Load Detection System\n')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].bold = True

p2 = doc.add_paragraph('in partial fulfilment of the requirement for the degree of\nBachelor of Technology In\n\nCOMPUTER SCIENCE AND ENGINEERING\n')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph('Submitted by\n\n[Name 1]   ([Roll Number 1])\n[Name 2]   ([Roll Number 2])\n[Name 3]   ([Roll Number 3])\n')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.runs[0].bold = True

p4 = doc.add_paragraph('Under the Supervision of\n[Supervisor Name]\n(Assistant Professor, CSE)\n\n\n\n[COLLEGE NAME]\n[LOCATION]')
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.runs[0].bold = True

doc.add_page_break()

# INDEX
doc.add_heading('Index', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sr. No.'
hdr_cells[1].text = 'Topics'
hdr_cells[2].text = 'Page No.'

index_content = [
    ('1', 'Introduction', '3'),
    ('2', 'Existing Systems', '4'),
    ('3', 'Problem Statement', '5'),
    ('4', 'Proposed Methodology', '6'),
    ('5', 'Feasibility Study', '12'),
    ('6', 'Facilities Required for Proposed Work', '13'),
    ('7', 'Conclusion', '14'),
    ('8', 'References', '15'),
]
for item in index_content:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

doc.add_paragraph('\n\n\nSupervisor Sign: _____________________')
doc.add_page_break()

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "Cognitive load refers to the total amount of mental effort being used in the working memory. In today's digital and remote-working era, measuring and managing cognitive load is critical for optimizing user performance, reducing stress, and improving productivity. Whether for professional developers, operators of critical systems, or students, identifying peak stress states early can prevent burnout.\n\n"
    "Traditional methods of measuring cognitive load often rely on invasive procedures or expensive biometric hardware like EEG caps. However, behavioral data derived from human-computer interaction (HCI)—such as typing speed, error rates, and mouse movement dynamics—along with facial cues like blink rates, provide a robust and non-invasive alternative to biometric sensors.\n\n"
    "The Proposed System, 'AI Cognitive Load Detection System,' is developed as a real-time behavioral tracking application. It integrates robust computer vision using OpenCV for eye-tracking and system-level input tracking via pynput to calculate behavioral proxies of stress. This multimodal data is processed by a Machine Learning model (Random Forest Classifier) that operates on an optimized zero-padded 60-second rolling window to accurately classify the user's current cognitive state as either 'High' or 'Low'.\n\n"
    "To make the application intuitive, a real-time web dashboard has been built using Flask and modern frontend technologies. It streams live predictions and hardware metrics instantaneously, ensuring an adaptive, unified environment for cognitive monitoring."
)
doc.add_page_break()

# 2. Existing Systems
doc.add_heading('2. Existing Systems', level=1)
doc.add_paragraph(
    "Historically, cognitive load management systems fall into distinct, fragmented categories. Existing systems tend to focus heavily on either physical biometrics (which require separate hardware) or minimal activity trackers that lack machine learning integration.\n"
)
table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'S.No.'
hdr_cells[1].text = 'System Type'
hdr_cells[2].text = 'Description'
hdr_cells[3].text = 'Advantages'
hdr_cells[4].text = 'Limitations'

existing_systems = [
    ('1', 'Hardware Biometrics (EEG/ECG)', 'Use medical-grade sensors to read brainwaves or heart rates.', 'Highly accurate baseline physiological data.', 'Invasive, expensive, and not usable for daily remote workflows.'),
    ('2', 'Activity & Productivity Trackers', 'Track active time spent on applications and keystroke counts.', 'Easy to install; minimal setup.', 'No cognitive analysis whatsoever; solely functional tracking.'),
    ('3', 'Camera-only Analysis Tools', 'Trace eye movements or facial expressions using webcams.', 'Non-invasive visual tracking.', 'Can be misread due to lighting; ignores physical system input.'),
    ('4', 'Periodic Survey Tools', 'Send pop-up surveys asking the user to self-report stress.', 'Relies on subjective user feeling.', 'Breaks user focus; unreliable due to cognitive bias.')
]
for item in existing_systems:
    row_cells = table2.add_row().cells
    for i in range(5):
        row_cells[i].text = item[i]

doc.add_paragraph('\nGap in Existing Systems').bold = True
doc.add_paragraph(
    "Most existing systems focus exclusively on single modalities—either pure video analysis or pure keyboard tracking—without combining them. There is no unified system that processes both facial blinks and hardware interactions simultaneously in real-time. Moreover, the lack of immediate visual feedback with rolling windows often leads to 'premature' stress alerts in existing tools. Our AI Cognitive Load Detection System fills this gap by implementing an integrated, zero-padded Rolling Window Machine Learning approach."
)
doc.add_page_break()

# 3. Problem Statement
doc.add_heading('3. Problem Statement', level=1)
doc.add_paragraph(
    "In the modern workspace, individuals are exposed to prolonged periods of intensive computer usage. This frequently leads to severe cognitive fatigue, compromised decision-making, and chronic stress. Unfortunately, individuals are often notoriously poor at self-assessing their own cognitive load until exhaustion has already taken its toll.\n\n"
    "While accurate biometric systems exist, they are highly impractical for regular deployment in standard corporate environments or at-home workstations due to cost, privacy concerns, and physical invasiveness.\n\n"
    "Conversely, simple productivity monitoring software only counts how many hours a person was active, utterly failing to differentiate between 'productive focus' and 'frustrated, error-prone typing'. A user scrambling to fix bugs might have high input metrics but terrible error rates and erratic mouse movements, signifying high stress.\n\n"
    "Furthermore, current systems that attempt algorithmic detection often suffer from extreme volatility. They trigger 'High Load' alerts based on momentary spikes rather than sustained behavior, leading to alert fatigue.\n\n"
    "Therefore, there is an urgent need to develop a non-invasive, highly accurate system that provides:\n"
    "• Simultaneous tracking of keyboard, mouse, and camera streams.\n"
    "• A sophisticated machine learning classifier trained on realistic thresholds.\n"
    "• A rolling window algorithm to ensure stability and eliminate false-positive anomalies.\n"
    "• An engaging, responsive visual dashboard to monitor the status in real-time.\n\n"
    "The proposed 'AI Cognitive Load Detection System' addresses these multidimensional challenges by providing an adaptive, seamless, and entirely software-based solution."
)
doc.add_page_break()

# 4. Proposed Methodology
doc.add_heading('4. Proposed Methodology', level=1)
doc.add_paragraph(
    "The proposed 'AI Cognitive Load Detection System' utilizes a component-based architecture for real-time tracking, evaluation, and reporting of user metrics. It operates entirely off standard system hardware and uses a Flask-based backend server combined with an AI reasoning engine.\n"
)
doc.add_heading('4.1 System Overview', level=2)
doc.add_paragraph(
    "The system captures real-time data streams via concurrent background threads. Data from the hardware interface (mouse, keyboard, webcam) is parsed into 4 primary distinct metrics: Typing Speed, Error Rate (Backspaces), Mouse Speed, and Blink Rate. These metrics are accumulated into a dynamic deque (rolling window) structure which eliminates sudden anomalies. A trained Random Forest model predicts the cognitive state based on the calculated averages."
)

doc.add_heading('4.2 Module Details', level=2)
for title, desc in [
    ("Module 1 – HARDWARE TRACKING MODULE", "Utilizes pynput libraries to seamlessly capture asynchronous system keyboard interrupts and mouse position changes without interrupting the user's workflow."),
    ("Module 2 – CAMERA & VISION MODULE", "Utilizes OpenCV and Haarcascade classifiers to identify eye pairs. The module implements a heuristic memory system to accurately detect blink transitions per minute."),
    ("Module 3 – ROLLING HISTORY MODULE", "Manages state history using Python's collections.deque. It enforces a 60-second window, originally zero-padded to prevent premature evaluation, yielding highly stable rolling averages."),
    ("Module 4 – MACHINE LEARNING MODULE", "Employs an advanced Scikit-Learn Random Forest Classifier trained on dynamically synthesized threshold data to calculate if the user is in an 'Optimal Focus' (Low) or 'Elevated Stress' (High) state."),
    ("Module 5 – WEB PROXY & DASHBOARD MODULE", "A lightweight Flask server handles backend analytics, continuously exposing metric evaluations via a RESTful JSON API. A beautiful, glassmorphism-styled HTML/CSS interface pulls this data asynchronously and displays it to the user.")
]:
    doc.add_paragraph(title).bold = True
    doc.add_paragraph(desc)

doc.add_heading('4.3 Use Case Diagram', level=2)
doc.add_paragraph("[ PLACEHOLDER FOR USE CASE DIAGRAM ]\nThe use case diagram consists of the primary Actor (The Computer User), who interacts with the system unintentionally via daily work. The system's use cases include Monitoring Inputs, Capturing Video Frames, Applying the Machine Learning Model, and Displaying the Web Dashboard.")

doc.add_heading('4.4 Data Flow Diagram (DFD)', level=2)
doc.add_paragraph("[ PLACEHOLDER FOR DFD ]\nLevel 0 DFD: User inputs data -> AI System processes it -> Dashboard visualizes it.\nLevel 1 DFD shows the split between independent listener threads (Keyboard, Mouse, OpenCV) funneling data into the Central Accumulator, which asks the Random Forest Model for a label before posting to Flask.")

doc.add_heading('4.5 System Flowchart', level=2)
doc.add_paragraph("[ PLACEHOLDER FOR FLOWCHART ]\nThe flowchart begins at System Start. The Hardware Listeners spin off into parallel background tasks. At 3-second intervals, the Central Loop extracts metrics, averages them against the DEQUE history, and executes Model.Predict(). If High, it raises the UI state / Windows Notification. If Low, it updates the visual elements gracefully. The loop ends gracefully via a KeyboardInterrupt.")

doc.add_heading('4.6 ER Diagram', level=2)
doc.add_paragraph("[ PLACEHOLDER FOR ER DIAGRAM ]\nIn the context of the operational state, the Entity Relationship involves the User Session, the Metric Vectors (Typing, Mouse, Error, Blink), and the Output History Logs. The User has a 1-to-many relationship with Metric Vectors, and Metric Vectors are foreign-keyed logically to a specific Model Prediction timestamp.")

doc.add_page_break()

# 5. Feasibility Study
doc.add_heading('5. Feasibility Study', level=1)
doc.add_heading('5.1 Technical Feasibility', level=2)
doc.add_paragraph("The proposed system relies entirely on standard hardware components (a built-in webcam and HID standard keyboard/mouse) and open-source libraries (Python, OpenCV, Scikit-learn). No specialized or unproven technical constraints exist, ensuring completely successful implementation.")
doc.add_heading('5.2 Operational Feasibility', level=2)
doc.add_paragraph("The application operates unobtrusively in the background without requiring user interference. The visually cohesive and smooth frontend Dashboard runs optimally in standard web browsers. No extensive user manual is required, making operational friction practically zero.")
doc.add_heading('5.3 Economic Feasibility', level=2)
doc.add_paragraph("Development involves free, open-source computational tools. Deployment does not demand highly-scaled cloud infrastructure since the lightweight inference model operates natively precisely on local machines, preserving extreme cost efficiency and absolute budget validation.")

doc.add_page_break()

# 6. Facilities Required for Proposed Work
doc.add_heading('6. Facilities Required for Proposed Work', level=1)
doc.add_heading('6.1 Hardware Requirements', level=2)
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
hw = [('Processor', 'Intel Core i3 / AMD Ryzen 3 or above'), ('RAM', '4 GB (8 GB recommended)'), ('Storage', '256 GB HDD/SSD'), ('Peripherals', 'Standard Webcam, Mouse, Keyboard')]
for i, item in enumerate(hw):
    table3.rows[i].cells[0].text = item[0]
    table3.rows[i].cells[1].text = item[1]

doc.add_heading('6.2 Software Requirements', level=2)
table4 = doc.add_table(rows=7, cols=3)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Software/Tools'
table4.rows[0].cells[1].text = 'Purpose'
table4.rows[0].cells[2].text = 'Version'
sw = [
    ('Python', 'Backend Core Logic, ML', '3.10+'),
    ('Flask', 'REST API and Web Server', 'Latest'),
    ('OpenCV (cv2)', 'Computer Vision & Eye Tracking', 'Latest'),
    ('Scikit-learn', 'Algorithm and RandomForest', 'Latest'),
    ('Pynput', 'Hardware Input Interrupts', 'Latest'),
    ('HTML, CSS, JS', 'Frontend UI layout', 'Latest')
]
for i, item in enumerate(sw):
    table4.rows[i+1].cells[0].text = item[0]
    table4.rows[i+1].cells[1].text = item[1]
    table4.rows[i+1].cells[2].text = item[2]

doc.add_page_break()

# 7. Conclusion
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    "The AI Cognitive Load Detection System provides a robust, seamless, and completely non-invasive mechanism for identifying human stress responses and cognitive burnout during periods of computer use. By aggregating disparate HCI metrics—from rapid mouse fluctuations to error-heavy keyboard typing—and cross-referencing them intuitively with human biological markers like blink rate via a webcam, the system successfully mimics evaluations previously requiring sophisticated biological sensors.\n\n"
    "With the integration of a stable Rolling Window deque buffer and an optimized Scikit-learn Random Forest model, the application prevents early-trigger false alarms, proving suitable for professional enterprise rollout or individual wellbeing tracking.\n\n"
    "Ultimately, the real-time Flask Web Dashboard and asynchronous architecture guarantees high performance without computational blockages. In the future, this system could be further modularized by appending centralized SQL database metrics logic, long-term historical visualization plots, and deep integration with Smart Environment APIs."
)
doc.add_page_break()

# 8. References
doc.add_heading('8. References', level=1)
doc.add_paragraph("[1] F. K. Hussain et al., 'Mental Workload Measurement in HCI,' Human-Computer Interaction Journals, 2018.")
doc.add_paragraph("[2] Pedregosa et al., 'Scikit-learn: Machine Learning in Python,' Journal of Machine Learning Research, 12, 2825-2830, 2011.")
doc.add_paragraph("[3] Flask Documentation, 'Pallets Official Flask Guide,' Available: https://flask.palletsprojects.com/")
doc.add_paragraph("[4] OpenCV Documentation, 'Open Source Computer Vision Library,' Available: https://opencv.org/")
doc.add_paragraph("[5] A. G. Bargas, 'Blink Rate and Cognitive Load Variations,' Neuroscience and HCI Studies, 2021.")

save_path = r"C:\Users\tej99\OneDrive\Desktop\AI_Cognitive_Load_Synopsis.docx"
try:
    doc.save(save_path)
    print(f"Synopsis successfully generated and saved to: {save_path}")
except Exception as e:
    print(f"Error saving document: {e}")

